"""
云服务文档编辑器模块
支持通过云服务API处理Office文档（docx、xlsx、pptx等）

当前支持：
1. Microsoft Graph API
2. 本地处理（fallback方案）
"""

import os
import json
from typing import Any, Dict, Optional, List
from abc import ABC, abstractmethod
from pathlib import Path
from settings.logger_manager import get_logger

logger = get_logger(__name__)


class CloudDocumentEditor(ABC):
    """云文档编辑器抽象基类"""
    
    @abstractmethod
    def create_document(self, content: str, output_path: str) -> bool:
        """创建新文档"""
        pass
    
    @abstractmethod
    def update_document(self, file_path: str, changes: Dict[str, Any]) -> bool:
        """更新文档内容"""
        pass
    
    @abstractmethod
    def format_document(self, file_path: str, formatting: Dict[str, Any]) -> bool:
        """格式化文档"""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """提取文档文本"""
        pass


class LocalDocumentEditor(CloudDocumentEditor):
    """本地文档编辑器（fallback方案）"""
    
    def __init__(self):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.enum.style import WD_STYLE_TYPE
            self.Document = Document
            self.Pt = Pt
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self.WD_LINE_SPACING = WD_LINE_SPACING
            self.WD_STYLE_TYPE = WD_STYLE_TYPE
            self.available = True
        except ImportError:
            logger.warning("python-docx 未安装，本地文档编辑功能不可用")
            self.available = False
    
    def create_document(self, content: str, output_path: str) -> bool:
        """创建新文档"""
        if not self.available:
            logger.error("python-docx 未安装")
            return False
        
        try:
            doc = self.Document()
            
            # 将内容按段落分割
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # 确保输出目录存在
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            doc.save(output_path)
            logger.info(f"文档创建成功: {output_path}")
            return True
        except Exception as e:
            logger.error(f"创建文档失败: {str(e)}")
            return False
    
    def update_document(self, file_path: str, changes: Dict[str, Any]) -> bool:
        """更新文档内容"""
        if not self.available:
            logger.error("python-docx 未安装")
            return False
        
        try:
            doc = self.Document(file_path)
            
            # 处理各种修改操作
            if 'replace_text' in changes:
                for old_text, new_text in changes['replace_text'].items():
                    for paragraph in doc.paragraphs:
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)
            
            if 'add_paragraph' in changes:
                for idx, para_content in enumerate(changes['add_paragraph']):
                    if 'position' in para_content:
                        pos = para_content['position']
                        if pos >= len(doc.paragraphs):
                            doc.add_paragraph(para_content['content'])
                        else:
                            doc.paragraphs[pos].insert_paragraph_before(para_content['content'])
                    else:
                        doc.add_paragraph(para_content['content'])
            
            if 'delete_paragraph' in changes:
                # 从后往前删除以避免索引问题
                positions = sorted(changes['delete_paragraph'], reverse=True)
                for pos in positions:
                    if 0 <= pos < len(doc.paragraphs):
                        p = doc.paragraphs[pos]
                        p._element.getparent().remove(p._element)
            
            if 'add_table' in changes:
                for table_data in changes['add_table']:
                    rows = table_data.get('rows', 3)
                    cols = table_data.get('cols', 3)
                    content = table_data.get('content', [])
                    table = doc.add_table(rows=rows, cols=cols)
                    for i, row in enumerate(content[:rows]):
                        for j, cell_content in enumerate(row[:cols]):
                            table.cell(i, j).text = str(cell_content)
            
            doc.save(file_path)
            logger.info(f"文档更新成功: {file_path}")
            return True
        except Exception as e:
            logger.error(f"更新文档失败: {str(e)}")
            return False
    
    def format_document(self, file_path: str, formatting: Dict[str, Any]) -> bool:
        """格式化文档"""
        if not self.available:
            logger.error("python-docx 未安装")
            return False
        
        try:
            doc = self.Document(file_path)
            
            # 应用格式设置
            if 'font_size' in formatting:
                font_size = self.Pt(formatting['font_size'])
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = font_size
            
            if 'font_name' in formatting:
                font_name = formatting['font_name']
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
            
            if 'font_color' in formatting:
                color = formatting['font_color']
                if color.startswith('#'):
                    r = int(color[1:3], 16)
                    g = int(color[3:5], 16)
                    b = int(color[5:7], 16)
                    rgb_color = self.RGBColor(r, g, b)
                else:
                    rgb_color = self.RGBColor(0, 0, 0)
                
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = rgb_color
            
            if 'alignment' in formatting:
                align_map = {
                    'left': self.WD_ALIGN_PARAGRAPH.LEFT,
                    'center': self.WD_ALIGN_PARAGRAPH.CENTER,
                    'right': self.WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': self.WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                align = align_map.get(formatting['alignment'].lower())
                if align:
                    for paragraph in doc.paragraphs:
                        paragraph.alignment = align
            
            if 'line_spacing' in formatting:
                spacing = formatting['line_spacing']
                for paragraph in doc.paragraphs:
                    paragraph.paragraph_format.line_spacing = spacing
            
            if 'margin' in formatting:
                margins = formatting['margin']
                sections = doc.sections
                for section in sections:
                    if 'top' in margins:
                        section.top_margin = self.Pt(margins['top'])
                    if 'bottom' in margins:
                        section.bottom_margin = self.Pt(margins['bottom'])
                    if 'left' in margins:
                        section.left_margin = self.Pt(margins['left'])
                    if 'right' in margins:
                        section.right_margin = self.Pt(margins['right'])
            
            doc.save(file_path)
            logger.info(f"文档格式化成功: {file_path}")
            return True
        except Exception as e:
            logger.error(f"格式化文档失败: {str(e)}")
            return False
    
    def extract_text(self, file_path: str) -> str:
        """提取文档文本"""
        if not self.available:
            logger.error("python-docx 未安装")
            return ""
        
        try:
            doc = self.Document(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return text
        except Exception as e:
            logger.error(f"提取文本失败: {str(e)}")
            return ""


class MicrosoftGraphEditor(CloudDocumentEditor):
    """Microsoft Graph API文档编辑器"""
    
    def __init__(self, client_id: str = None, client_secret: str = None, tenant_id: str = None):
        self.client_id = client_id or os.getenv('AZURE_CLIENT_ID')
        self.client_secret = client_secret or os.getenv('AZURE_CLIENT_SECRET')
        self.tenant_id = tenant_id or os.getenv('AZURE_TENANT_ID')
        self.access_token = None
        self.available = self._check_config()
    
    def _check_config(self) -> bool:
        """检查配置是否完整"""
        if not self.client_id or not self.client_secret or not self.tenant_id:
            logger.warning("Microsoft Graph API 配置不完整，将使用本地编辑器")
            return False
        return True
    
    def _get_access_token(self) -> bool:
        """获取访问令牌"""
        if not self.available:
            return False
        
        try:
            import requests
            
            url = f"https://login.microsoftonline.com/{self.tenant_id}/oauth2/v2.0/token"
            data = {
                'client_id': self.client_id,
                'client_secret': self.client_secret,
                'grant_type': 'client_credentials',
                'scope': 'https://graph.microsoft.com/.default'
            }
            
            response = requests.post(url, data=data)
            response.raise_for_status()
            self.access_token = response.json().get('access_token')
            return True
        except Exception as e:
            logger.error(f"获取Microsoft Graph API令牌失败: {str(e)}")
            return False
    
    def create_document(self, content: str, output_path: str) -> bool:
        """创建新文档（使用OneDrive）"""
        if not self.available:
            return self._fallback_create(content, output_path)
        
        if not self.access_token and not self._get_access_token():
            return self._fallback_create(content, output_path)
        
        try:
            import requests
            
            # 创建文档内容
            document_content = {
                "name": os.path.basename(output_path),
                "content": content,
                "type": "docx"
            }
            
            url = "https://graph.microsoft.com/v1.0/me/drive/root:/Documents/{filename}:/content"
            url = url.format(filename=os.path.basename(output_path))
            
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
            # 先创建本地文件，然后上传
            local_editor = LocalDocumentEditor()
            temp_path = f"{output_path}.temp"
            if local_editor.create_document(content, temp_path):
                with open(temp_path, 'rb') as f:
                    response = requests.put(url, headers=headers, data=f)
                    response.raise_for_status()
                
                os.remove(temp_path)
                logger.info(f"文档通过Microsoft Graph API创建成功")
                return True
            
            return False
        except Exception as e:
            logger.error(f"Microsoft Graph API创建文档失败: {str(e)}")
            return self._fallback_create(content, output_path)
    
    def update_document(self, file_path: str, changes: Dict[str, Any]) -> bool:
        """更新文档内容"""
        if not self.available:
            return self._fallback_update(file_path, changes)
        
        if not self.access_token and not self._get_access_token():
            return self._fallback_update(file_path, changes)
        
        try:
            # 对于复杂修改，先下载到本地，修改后再上传
            local_editor = LocalDocumentEditor()
            if local_editor.update_document(file_path, changes):
                # 上传更新后的文件到OneDrive
                self._upload_to_onedrive(file_path)
                return True
            return False
        except Exception as e:
            logger.error(f"Microsoft Graph API更新文档失败: {str(e)}")
            return self._fallback_update(file_path, changes)
    
    def format_document(self, file_path: str, formatting: Dict[str, Any]) -> bool:
        """格式化文档"""
        if not self.available:
            return self._fallback_format(file_path, formatting)
        
        if not self.access_token and not self._get_access_token():
            return self._fallback_format(file_path, formatting)
        
        try:
            local_editor = LocalDocumentEditor()
            if local_editor.format_document(file_path, formatting):
                self._upload_to_onedrive(file_path)
                return True
            return False
        except Exception as e:
            logger.error(f"Microsoft Graph API格式化文档失败: {str(e)}")
            return self._fallback_format(file_path, formatting)
    
    def extract_text(self, file_path: str) -> str:
        """提取文档文本"""
        if not self.available:
            return self._fallback_extract(file_path)
        
        try:
            local_editor = LocalDocumentEditor()
            return local_editor.extract_text(file_path)
        except Exception as e:
            logger.error(f"Microsoft Graph API提取文本失败: {str(e)}")
            return self._fallback_extract(file_path)
    
    def _upload_to_onedrive(self, file_path: str):
        """上传文件到OneDrive"""
        try:
            import requests
            
            filename = os.path.basename(file_path)
            url = f"https://graph.microsoft.com/v1.0/me/drive/root:/Documents/{filename}:/content"
            
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
            with open(file_path, 'rb') as f:
                response = requests.put(url, headers=headers, data=f)
                response.raise_for_status()
            
            logger.info(f"文件已上传到OneDrive: {filename}")
        except Exception as e:
            logger.error(f"上传到OneDrive失败: {str(e)}")
    
    def _fallback_create(self, content: str, output_path: str) -> bool:
        """fallback到本地编辑器"""
        logger.info("使用本地编辑器创建文档")
        return LocalDocumentEditor().create_document(content, output_path)
    
    def _fallback_update(self, file_path: str, changes: Dict[str, Any]) -> bool:
        """fallback到本地编辑器"""
        logger.info("使用本地编辑器更新文档")
        return LocalDocumentEditor().update_document(file_path, changes)
    
    def _fallback_format(self, file_path: str, formatting: Dict[str, Any]) -> bool:
        """fallback到本地编辑器"""
        logger.info("使用本地编辑器格式化文档")
        return LocalDocumentEditor().format_document(file_path, formatting)
    
    def _fallback_extract(self, file_path: str) -> str:
        """fallback到本地编辑器"""
        logger.info("使用本地编辑器提取文本")
        return LocalDocumentEditor().extract_text(file_path)


class DocumentEditorFactory:
    """文档编辑器工厂类"""
    
    @staticmethod
    def create_editor(cloud_provider: str = 'local') -> CloudDocumentEditor:
        """创建文档编辑器实例"""
        if cloud_provider.lower() == 'microsoft' or cloud_provider.lower() == 'azure':
            return MicrosoftGraphEditor()
        elif cloud_provider.lower() == 'google':
            logger.warning("Google Drive API编辑器尚未实现，使用本地编辑器")
            return LocalDocumentEditor()
        elif cloud_provider.lower() == 'aliyun':
            logger.warning("阿里云文档API编辑器尚未实现，使用本地编辑器")
            return LocalDocumentEditor()
        elif cloud_provider.lower() == 'tencent':
            logger.warning("腾讯云文档API编辑器尚未实现，使用本地编辑器")
            return LocalDocumentEditor()
        else:
            return LocalDocumentEditor()


# 全局编辑器实例
document_editor = DocumentEditorFactory.create_editor()


if __name__ == "__main__":
    # 测试代码
    editor = DocumentEditorFactory.create_editor('local')
    
    # 测试创建文档
    test_content = """欢迎使用云文档编辑器

这是一个测试文档，用于验证docx文档的创建和修改功能。

功能特点：
- 支持创建新文档
- 支持更新文档内容
- 支持格式化文档
- 支持提取文本内容
"""
    
    test_file = "test_output.docx"
    
    # 创建文档
    print("测试创建文档...")
    result = editor.create_document(test_content, test_file)
    print(f"创建结果: {result}")
    
    # 测试提取文本
    print("\n测试提取文本...")
    text = editor.extract_text(test_file)
    print(f"提取的文本长度: {len(text)}")
    
    # 测试更新文档
    print("\n测试更新文档...")
    changes = {
        'replace_text': {'测试文档': '示例文档'},
        'add_paragraph': [{'content': '这是新增的段落'}]
    }
    result = editor.update_document(test_file, changes)
    print(f"更新结果: {result}")
    
    # 测试格式化
    print("\n测试格式化文档...")
    formatting = {
        'font_size': 12,
        'font_name': '微软雅黑',
        'alignment': 'justify',
        'line_spacing': 1.5,
        'margin': {'top': 25.4, 'bottom': 25.4, 'left': 25.4, 'right': 25.4}
    }
    result = editor.format_document(test_file, formatting)
    print(f"格式化结果: {result}")
    
    print("\n所有测试完成！")